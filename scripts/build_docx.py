#!/usr/bin/env python3
"""
build_docx.py
Converte Markdown estruturado -> DOCX profissional (python-docx)
"""

import argparse
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, str(Path(__file__).resolve().parent))
from utils.md_parser import parse_markdown, BlockType, parse_inline


def _set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_paragraph_shading(paragraph, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    paragraph.paragraph_format.element.get_or_add_pPr().append(shading)


def _set_paragraph_border(paragraph, color, sz):
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ["top", "left", "bottom", "right"]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), str(sz))
        elem.set(qn("w:space"), "1")
        elem.set(qn("w:color"), color)
        pBdr.append(elem)
    pPr.append(pBdr)


def text_color_str(bg_hex):
    """Return a darker shade for border from a hex background color."""
    mapping = {"E3F2FD": "1565C0", "FFF8E1": "F57F17", "E8F5E9": "2E7D32",
               "FFEBEE": "C62828", "ECEFF1": "546E7A", "F5F5F5": "1A237E"}
    return mapping.get(bg_hex.upper(), "333333")


class DocxBuilder:
    def __init__(self):
        self.doc = Document()

    def setup_styles(self):
        style = self.doc.styles["Normal"]
        font = style.font
        font.name = "Georgia"
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x21, 0x21, 0x21)
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.6

        for level, size, color in [
            (1, 24, "1A237E"),
            (2, 18, "1A237E"),
            (3, 14, "0D47A1"),
            (4, 12, "212121"),
        ]:
            hs = self.doc.styles[f"Heading {level}"]
            hs.font.name = "Segoe UI"
            hs.font.size = Pt(size)
            hs.font.bold = True
            hs.font.color.rgb = RGBColor(*bytes.fromhex(color))
            hs.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            hs.paragraph_format.space_after = Pt(6)

    def add_cover_page(self, title: str, subtitle: str = "", author: str = ""):
        for _ in range(6):
            self.doc.add_paragraph()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(subtitle)
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x00, 0xBF, 0xA5)

        self.doc.add_paragraph()

        if author:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(author)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

        self.doc.add_page_break()

    def add_toc(self):
        self.doc.add_heading("Sumario", level=1)
        p = self.doc.add_paragraph()
        run = p.add_run(
            "(Atualize o sumario no Word: clique com botao direito -> Atualizar Campo)"
        )
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        self.doc.add_page_break()

    def add_section_break(self):
        self.doc.add_page_break()

    def add_heading(self, text: str, level: int):
        self.doc.add_heading(text, level=level)

    def add_paragraph(self, text: str):
        segments = parse_inline(text)

        if not segments:
            return

        p = self.doc.add_paragraph()
        for seg in segments:
            run = p.add_run(seg.text)
            if seg.bold:
                run.bold = True
            if seg.italic:
                run.italic = True
            if seg.code:
                run.font.name = "Cascadia Code"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)

    def add_code_block(self, code: str, language: str = ""):
        p = self.doc.add_paragraph()
        code = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', code)
        run = p.add_run(code)
        run.font.name = "Cascadia Code"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xE0, 0xE0, 0xE0)
        pf = p.paragraph_format
        pf.left_indent = Cm(1)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        _set_paragraph_shading(p, "263238")
        _set_paragraph_border(p, "E0E0E0", 1)

    def add_mermaid_block(self, code: str):
        p = self.doc.add_paragraph()
        run = p.add_run("[Diagrama]")
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
        run.font.size = Pt(9)
        pf = p.paragraph_format
        pf.space_before = Pt(6)

        p2 = self.doc.add_paragraph()
        run2 = p2.add_run(code)
        run2.font.name = "Cascadia Code"
        run2.font.size = Pt(8)
        run2.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
        pf2 = p2.paragraph_format
        pf2.left_indent = Cm(1)
        pf2.space_after = Pt(6)
        _set_paragraph_shading(p2, "E3F2FD")

    def add_list(self, items: list, ordered: bool = False):
        for item in items:
            self.doc.add_paragraph(item, style="List Number" if ordered else "List Bullet")

    def _add_callout_paragraph(self, text: str, label: str, bg_color: str, text_color, font_style=None):
        p = self.doc.add_paragraph()

        # Add label run
        if label:
            run = p.add_run(label)
            run.font.color.rgb = text_color
            if font_style == "italic":
                run.font.italic = True
            elif font_style == "bold":
                run.font.bold = True

        # Add content with inline formatting
        segments = parse_inline(text)
        for seg in segments:
            run = p.add_run(seg.text)
            run.font.color.rgb = text_color
            if seg.bold:
                run.bold = True
            if seg.italic:
                run.italic = True
            if seg.code:
                run.font.name = "Cascadia Code"
                run.font.size = Pt(8)
            elif font_style == "italic":
                run.font.italic = True
            elif font_style == "bold":
                run.font.bold = True

        # Add left border stripe for container feel
        pPr = p.paragraph_format.element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "12")  # 6pt border
        left.set(qn("w:space"), "4")
        # Extract border color from bg_color (convert light to darker shade)
        left.set(qn("w:color"), text_color_str(bg_color))
        pBdr.append(left)
        pPr.append(pBdr)
        pf = p.paragraph_format
        pf.left_indent = Cm(1)
        _set_paragraph_shading(p, bg_color)

    def add_tip(self, text: str):
        self._add_callout_paragraph(
            text, "Dica: ",
            "E3F2FD", RGBColor(0x0D, 0x47, 0xA1),
            font_style="italic"
        )

    def add_warning(self, text: str):
        self._add_callout_paragraph(
            text, "Atencao: ",
            "FFF8E1", RGBColor(0xF5, 0x7F, 0x17),
            font_style="bold"
        )

    def add_definition(self, text: str):
        self._add_callout_paragraph(
            text, "",
            "E8F5E9", RGBColor(0x2E, 0x7D, 0x32)
        )

    def add_info(self, text: str):
        self._add_callout_paragraph(
            text, "",
            "ECEFF1", RGBColor(0x42, 0x42, 0x42)
        )

    def add_recap(self, text: str):
        self._add_callout_paragraph(
            text, "Recapitulando: ",
            "F5F5F5", RGBColor(0x1A, 0x23, 0x7E),
            font_style="bold"
        )

    def add_table(self, table_text: str):
        rows = table_text.strip().split("\n")
        data = [row.strip(" |").split("|") for row in rows if row.strip()]
        if len(data) < 2:
            return
        # Skip separator row
        data = [row for row in data if not all(cell.strip().startswith("-") for cell in row)]
        if not data:
            return

        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Light Grid Accent 1"
        for i, row_data in enumerate(data):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text.strip()
        # Style header row
        for j, cell in enumerate(table.rows[0].cells):
            _set_cell_shading(cell, "1A237E")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def add_hr(self):
        p = self.doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "00BFA5")
        pBdr.append(bottom)
        p.paragraph_format.element.get_or_add_pPr().append(pBdr)

    def convert_markdown(self, md_text: str, title: str = ""):
        doc = parse_markdown(md_text)

        for block in doc.blocks:
            if block.type == BlockType.HEADING:
                self.add_heading(block.content, block.level)
            elif block.type == BlockType.PARAGRAPH:
                self.add_paragraph(block.content)
            elif block.type == BlockType.CODE:
                self.add_code_block(block.content, block.language)
            elif block.type == BlockType.MERMAID:
                self.add_mermaid_block(block.content)
            elif block.type == BlockType.LIST:
                has_numbers = any(c.isdigit() for c in str(block.items[:1]))
                self.add_list(block.items, ordered=has_numbers)
            elif block.type == BlockType.TABLE:
                self.add_table(block.content)
            elif block.type == BlockType.TIP:
                self.add_tip(block.content)
            elif block.type == BlockType.WARNING:
                self.add_warning(block.content)
            elif block.type == BlockType.BLOCKQUOTE:
                p = self.doc.add_paragraph()
                run = p.add_run(block.content)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
                run.font.size = Pt(10)
                pf = p.paragraph_format
                pf.left_indent = Cm(0.5)
                # Green left border
                pPr = p.paragraph_format.element.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                left = OxmlElement("w:left")
                left.set(qn("w:val"), "single")
                left.set(qn("w:sz"), "8")
                left.set(qn("w:space"), "4")
                left.set(qn("w:color"), "00BFA5")
                pBdr.append(left)
                pPr.append(pBdr)
            elif block.type == BlockType.DEFINICAO:
                self.add_definition(block.content)
            elif block.type == BlockType.INFO:
                self.add_info(block.content)
            elif block.type == BlockType.RECAPITULANDO:
                self._add_callout_paragraph(
                    block.content, "Recapitulando: ",
                    "F5F5F5", RGBColor(0x1A, 0x23, 0x7E),
                    font_style="bold"
                )
            elif block.type == BlockType.HR:
                self.add_hr()

    def add_header_footer(self, title: str):
        section = self.doc.sections[0]
        header = section.header
        hp = header.paragraphs[0]
        hp.text = title
        hp.style.font.size = Pt(8)
        hp.style.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        run2 = fp.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = " PAGE "
        run2._r.append(instrText)
        run3 = fp.add_run()
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")
        run3._r.append(fldChar3)

    def save(self, path: Path):
        self.doc.core_properties.title = self._title if hasattr(self, "_title") and self._title else path.stem
        self.doc.core_properties.author = self._author if hasattr(self, "_author") and self._author else "AI Software Engineering Academy"
        self.doc.core_properties.category = "Technology, Artificial Intelligence, Software Development"
        self.doc.core_properties.comments = "ISBN: 978-65-992345-00-0 | CC BY-NC-SA 4.0 | WCAG AA | pt-BR"
        self.doc.core_properties.keywords = "IA, inteligencia artificial, programacao, desenvolvimento, automacao, CI/CD"
        self.doc.core_properties.language = "pt-BR"
        self.doc.save(str(path))
        print(f"  OK DOCX salvo: {path}")


def main():
    parser = argparse.ArgumentParser(description="Build DOCX from Markdown")
    parser.add_argument("--input", required=True, help="Arquivo Markdown de entrada")
    parser.add_argument("--output", required=True, help="Arquivo DOCX de saida")
    parser.add_argument("--title", default="", help="Titulo do documento")
    parser.add_argument("--subtitle", default="", help="Subtitulo")
    parser.add_argument("--author", default="AI Software Engineering Academy", help="Autor")
    parser.add_argument("--cover", action="store_true", help="Incluir pagina de capa")
    parser.add_argument("--toc", action="store_true", help="Incluir sumario")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"ERRO: Arquivo nao encontrado: {input_path}")
        sys.exit(1)

    md_text = input_path.read_text(encoding="utf-8")
    title = args.title or input_path.stem

    builder = DocxBuilder()
    builder.setup_styles()

    if args.cover:
        builder.add_cover_page(title, args.subtitle, args.author)
    if args.toc:
        builder.add_toc()

    print(f"  Convertendo {input_path.name} -> DOCX...")
    builder.convert_markdown(md_text, title)
    builder.add_header_footer(title)

    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    builder.save(output_path)


if __name__ == "__main__":
    main()
